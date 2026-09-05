"""
boss_agent.memory
=================
Resume ingestion, text extraction, LLM structured profile generation, and local memory cache.
"""

import json
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

import yaml
from rich.console import Console

from droid_agent_core.llm import LLMDecisionClient, OpenAIChatClient

console = Console()


@dataclass
class StructuredCandidateProfile:
    """Structured memory profile representing candidate background and key strengths."""

    name: str = "求职者"
    years_of_experience: int = 0
    education: list[dict[str, str]] = field(default_factory=list)
    core_skills: list[str] = field(default_factory=list)
    work_experiences: list[dict[str, Any]] = field(default_factory=list)
    projects: list[dict[str, Any]] = field(default_factory=list)
    project_highlights: list[dict[str, str]] = field(default_factory=list)
    target_positions: list[str] = field(default_factory=list)
    raw_summary: str = ""
    raw_resume_text: str = ""

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "StructuredCandidateProfile":
        raw_skills = data.get("core_skills") or []
        normalized_skills: list[str] = []
        if isinstance(raw_skills, dict):
            for k, v in raw_skills.items():
                if isinstance(v, list):
                    normalized_skills.append(f"{k}: {', '.join(str(x) for x in v)}")
                else:
                    normalized_skills.append(f"{k}: {v}")
        elif isinstance(raw_skills, list):
            for item in raw_skills:
                if isinstance(item, dict):
                    for k, v in item.items():
                        if isinstance(v, list):
                            normalized_skills.append(f"{k}: {', '.join(str(x) for x in v)}")
                        else:
                            normalized_skills.append(f"{k}: {v}")
                else:
                    normalized_skills.append(str(item))

        raw_work = data.get("work_experiences") or []
        raw_projects = data.get("projects") or []
        raw_highlights = data.get("project_highlights") or []
        if not raw_projects and raw_highlights:
            raw_projects = [
                {
                    "name": p.get("name", ""),
                    "description": p.get("description", ""),
                    "role": p.get("role", ""),
                    "tech_stack": p.get("tech_stack", []),
                    "achievements": p.get("achievements", ""),
                    "raw_details": p.get("raw_details", ""),
                }
                for p in raw_highlights
            ]
        if not raw_highlights and raw_projects:
            raw_highlights = [
                {
                    "name": p.get("name", ""),
                    "description": p.get("description", "") or p.get("achievements", ""),
                }
                for p in raw_projects
            ]

        return cls(
            name=data.get("name") or "求职者",
            years_of_experience=int(data.get("years_of_experience") or 0),
            education=data.get("education") or [],
            core_skills=normalized_skills,
            work_experiences=raw_work,
            projects=raw_projects,
            project_highlights=raw_highlights,
            target_positions=data.get("target_positions") or [],
            raw_summary=data.get("raw_summary") or "",
            raw_resume_text=data.get("raw_resume_text") or "",
        )

    def format_for_prompt(self) -> str:
        """Format the profile into an unabbreviated, detail-rich block for LLM prompts."""
        edu_str = "; ".join(
            f"{e.get('school', '')} ({e.get('degree', '')} - {e.get('major', '')})"
            for e in self.education
        )
        skills_str = "\n".join(f"- {s}" for s in self.core_skills) if self.core_skills else "未注明"
        targets_str = ", ".join(self.target_positions)

        work_items = []
        for w in self.work_experiences:
            comp = w.get("company", "")
            role = w.get("role", "")
            start = w.get("start_date", "")
            end = w.get("end_date", "")
            time_span = f" ({start} ~ {end})" if start or end else ""
            dept = f" | 部门: {w.get('department')}" if w.get("department") else ""
            header = f"- 【{comp}】{role}{time_span}{dept}"
            details = []
            if w.get("responsibilities"):
                details.append(f"  工作职责: {w.get('responsibilities')}")
            if w.get("achievements"):
                details.append(f"  核心业绩与量化成果: {w.get('achievements')}")
            if w.get("raw_details"):
                details.append(f"  详细履历: {w.get('raw_details')}")
            work_items.append(header + ("\n" + "\n".join(details) if details else ""))
        work_str = "\n".join(work_items) if work_items else "未注明"

        active_projects = self.projects if self.projects else self.project_highlights
        proj_items = []
        for p in active_projects:
            p_name = p.get("name", "")
            p_role = f" (角色: {p.get('role')})" if p.get("role") else ""
            start = p.get("start_date", "")
            end = p.get("end_date", "")
            p_time = f" [{start} ~ {end}]" if start or end else ""
            stack_val = p.get("tech_stack")
            stack_str = (
                f" | 技术栈: {', '.join(stack_val) if isinstance(stack_val, list) else stack_val}"
                if stack_val
                else ""
            )
            header = f"- 【{p_name}】{p_role}{p_time}{stack_str}"
            p_details = []
            if p.get("description"):
                p_details.append(f"  项目背景与架构: {p.get('description')}")
            if p.get("achievements"):
                p_details.append(f"  核心贡献与指标成果: {p.get('achievements')}")
            if p.get("raw_details"):
                p_details.append(f"  技术攻坚细节: {p.get('raw_details')}")
            proj_items.append(header + ("\n" + "\n".join(p_details) if p_details else ""))
        projects_str = "\n".join(proj_items) if proj_items else "未注明"

        ground_truth = (
            f"\n\n[原始简历无损语料 (Ground Truth 参考)]\n{self.raw_resume_text.strip()}"
            if self.raw_resume_text and self.raw_resume_text.strip()
            else ""
        )

        return (
            f"姓名: {self.name}\n"
            f"工作经验: {self.years_of_experience}年\n"
            f"教育背景: {edu_str or '未注明'}\n"
            f"核心技能栈:\n{skills_str}\n"
            f"期望职位: {targets_str or '未注明'}\n"
            f"工作经历 (无损完整履历):\n{work_str}\n"
            f"项目经历 (完整架构与指标):\n{projects_str}\n"
            f"个人总结与背景优势: {self.raw_summary or '未注明'}"
            f"{ground_truth}"
        )


class ResumeTextExtractor:
    """Extracts raw text content from local resume files (.pdf, .docx, .txt, .md, .json)."""

    def extract_text(self, file_path: str | Path) -> str:
        path = Path(file_path)
        if not path.is_file():
            raise FileNotFoundError(f"Resume file not found at: {file_path}")

        ext = path.suffix.lower()
        if ext in [".txt", ".md", ".json"]:
            return self._read_text_file(path)
        elif ext == ".pdf":
            return self._read_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self._read_docx(path)
        else:
            # Fallback to plain text read
            return self._read_text_file(path)

    def _read_text_file(self, path: Path) -> str:
        for encoding in ["utf-8", "gb18030", "gbk", "utf-16"]:
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_bytes().decode("utf-8", errors="replace")

    def _read_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            return "\n\n".join(pages_text)
        except Exception as e:
            console.print(
                f"[yellow]⚠️  pypdf extraction failed ({e}), falling back to text read[/yellow]"
            )
            return self._read_text_file(path)

    def _read_docx(self, path: Path) -> str:
        try:
            import docx

            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n".join(paragraphs)
        except Exception as e:
            console.print(
                f"[yellow]⚠️  python-docx extraction failed ({e}), falling back to text read[/yellow]"
            )
            return self._read_text_file(path)


class ResumeMemoryManager:
    """Manages parsing, structuring, and local caching of candidate profile memory."""

    DEFAULT_MEMORY_PATH = Path("config/candidate_memory.json")

    def __init__(
        self,
        llm_client: LLMDecisionClient | None = None,
        memory_file_path: str | Path | None = None,
        candidate_config_path: str | Path | None = None,
    ):
        self.llm_client = llm_client or OpenAIChatClient()
        self.extractor = ResumeTextExtractor()

        # Load candidate config if available
        self.candidate_config = self._load_candidate_config(candidate_config_path)

        configured_memory = (
            memory_file_path or self.candidate_config.get("memory_path") or self.DEFAULT_MEMORY_PATH
        )
        self.memory_path = Path(configured_memory)
        self.configured_resume_path = self.candidate_config.get("resume_path")

    @staticmethod
    def _load_candidate_config(config_path: str | Path | None = None) -> dict[str, Any]:
        search_paths: list[Path] = []
        if config_path:
            search_paths.append(Path(config_path))
        else:
            search_paths.extend(
                [
                    Path("config/candidate.local.yaml"),
                    Path("config/candidate.local.json"),
                    Path("config/candidate.yaml"),
                    Path("config/candidate.json"),
                ]
            )

        for p in search_paths:
            if p.is_file():
                try:
                    content = p.read_text(encoding="utf-8")
                    if p.suffix in [".yaml", ".yml"]:
                        loaded = yaml.safe_load(content) or {}
                    else:
                        loaded = json.loads(content) or {}
                    if isinstance(loaded, dict):
                        return loaded
                except Exception:
                    pass
        return {}

    def has_memory_file(self) -> bool:
        """Return True if candidate memory profile file exists and is non-empty."""
        return self.memory_path.is_file() and self.memory_path.stat().st_size > 0

    def load_cached_memory(self) -> StructuredCandidateProfile | None:
        """Load memory profile from PocketBase database (single source of truth) with file fallback."""
        # 1. Primary: load from PocketBase database (HTTP or SQLite fallback)
        try:
            import asyncio

            from boss_agent.broker import PocketBaseBroker

            broker = PocketBaseBroker()
            try:
                asyncio.get_running_loop()
                import concurrent.futures

                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                    data = pool.submit(asyncio.run, broker.get_candidate_profile()).result(
                        timeout=3.0
                    )
            except RuntimeError:
                data = asyncio.run(broker.get_candidate_profile())

            if data and (
                data.get("name")
                or data.get("core_skills")
                or data.get("work_experiences")
                or data.get("projects")
            ):
                return StructuredCandidateProfile.from_dict(data)
        except Exception as e:
            console.print(f"[dim]Note: Database candidate profile check: {e}[/dim]")

        # 2. File fallback if memory_path exists
        if self.has_memory_file():
            try:
                content = self.memory_path.read_text(encoding="utf-8")
                data = json.loads(content)
                return StructuredCandidateProfile.from_dict(data)
            except Exception as e:
                console.print(
                    f"[yellow]⚠️  Failed to read cached memory from {self.memory_path}: {e}[/yellow]"
                )
        return None

    def generate_and_save_memory(self, resume_path: str | Path) -> StructuredCandidateProfile:
        """Extract text from resume file, call LLM to parse into unabbreviated schema, and save to database."""
        console.print(f"📄 [bold cyan]Parsing resume file:[/bold cyan] {resume_path}...")
        raw_text = self.extractor.extract_text(resume_path)
        if not raw_text.strip():
            raise ValueError(f"Extracted resume text from {resume_path} was empty.")

        console.print(
            "🧠 [bold cyan]Structuring candidate profile via LLM (unabbreviated)...[/bold cyan]"
        )
        prompt = (
            "请全面、无损、完整地解析以下求职者原始简历文本，并提取出结构化信息：\n\n"
            f"[简历文本内容]\n{raw_text}\n\n"
            "请严格以标准 JSON 格式输出以下结构。严禁进行摘要删减，必须全量保留所有工作经历与项目履历中的职责、业务场景、技术选型与量化成果：\n"
            "{\n"
            '  "name": "姓名",\n'
            '  "years_of_experience": 经验年限(整数),\n'
            '  "education": [{"school": "学校", "degree": "学历", "major": "专业", "start_date": "开始时间", "end_date": "结束时间"}],\n'
            '  "core_skills": ["分类1: 技能列表", "分类2: 技能列表"],\n'
            '  "work_experiences": [\n'
            "    {\n"
            '      "company": "公司名称",\n'
            '      "role": "职位/角色",\n'
            '      "start_date": "起止年或年月",\n'
            '      "end_date": "结束年月或至今",\n'
            '      "department": "所属部门/业务线",\n'
            '      "responsibilities": "完整详细的工作职责描述与主导内容",\n'
            '      "achievements": "核心量化业绩、技术攻坚指标与业务成果",\n'
            '      "raw_details": "该段经历的完整原文及补充细节"\n'
            "    }\n"
            "  ],\n"
            '  "projects": [\n'
            "    {\n"
            '      "name": "项目名称",\n'
            '      "role": "担任角色",\n'
            '      "start_date": "开始时间",\n'
            '      "end_date": "结束时间",\n'
            '      "tech_stack": ["技术1", "技术2"],\n'
            '      "description": "完整项目背景、技术架构与负责模块",\n'
            '      "achievements": "项目可量化成果、指标突破与关键产出",\n'
            '      "raw_details": "技术攻坚与落地细节"\n'
            "    }\n"
            "  ],\n"
            '  "target_positions": ["期望职位1", "期望职位2"],\n'
            '  "raw_summary": "全面总结的核心个人背景亮点与技术优势概述"\n'
            "}\n\n"
            "注意：必须输出标准严格合法的 JSON。所有经历必须完整提取，严禁做概括删减；字符串内严禁未转义的双引号（若需引用请用中文书名号《》或单引号）。"
        )
        messages = [
            {
                "role": "system",
                "content": "You are a professional HR assistant specializing in parsing candidate resumes into structured JSON with 100% fidelity. Always output valid, complete JSON.",
            },
            {"role": "user", "content": prompt},
        ]

        extra_payload = None
        if (
            "minimax" in getattr(self.llm_client.config, "base_url", "").lower()
            or "minimax" in getattr(self.llm_client.config, "model", "").lower()
        ):
            extra_payload = {"thinking": {"type": "disabled"}}

        result_dict = self.llm_client.chat_completion_json(
            messages,
            max_tokens=16384,
            extra_payload=extra_payload,
        )
        result_dict["raw_resume_text"] = raw_text
        profile = StructuredCandidateProfile.from_dict(result_dict)

        self.save_memory_profile(profile)
        return profile

    def save_memory_profile(self, profile: StructuredCandidateProfile) -> None:
        """Save candidate profile to PocketBase database with local file fallback."""
        # 1. Primary: Save to PocketBase database (HTTP or SQLite fallback)
        try:
            import asyncio

            from boss_agent.broker import PocketBaseBroker

            broker = PocketBaseBroker()
            try:
                asyncio.get_running_loop()
                import concurrent.futures

                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                    pool.submit(
                        asyncio.run, broker.save_candidate_profile(profile.to_dict())
                    ).result(timeout=5.0)
            except RuntimeError:
                asyncio.run(broker.save_candidate_profile(profile.to_dict()))
            console.print(
                "✅ [bold green]Structured candidate profile saved to PocketBase database.[/bold green]"
            )
        except Exception as e:
            console.print(f"[yellow]⚠️  Failed to save profile to database: {e}[/yellow]")

        # 2. Save to local file if path is specified
        try:
            self.memory_path.parent.mkdir(parents=True, exist_ok=True)
            self.memory_path.write_text(
                json.dumps(profile.to_dict(), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except Exception:
            pass

    def load_memory(
        self,
        force_refresh: bool = False,
        resume_file: str | Path | None = None,
    ) -> StructuredCandidateProfile:
        """Load candidate profile memory idempotently or regenerate from resume."""
        target_resume = resume_file or self.configured_resume_path

        if not force_refresh and not target_resume:
            cached = self.load_cached_memory()
            if cached is not None:
                console.print(
                    f"💾 [dim]Loaded existing candidate memory from {self.memory_path}[/dim]"
                )
                return cached

        if target_resume and (force_refresh or not self.has_memory_file()):
            return self.generate_and_save_memory(target_resume)

        if not force_refresh:
            cached = self.load_cached_memory()
            if cached is not None:
                console.print(
                    f"💾 [dim]Loaded existing candidate memory from {self.memory_path}[/dim]"
                )
                return cached

        # Look for default resumes directory
        default_resumes_dir = Path("resumes")
        if default_resumes_dir.is_dir():
            candidates = list(default_resumes_dir.glob("*.*"))
            valid_candidates = [
                c
                for c in candidates
                if c.suffix.lower() in [".pdf", ".docx", ".doc", ".txt", ".md"]
            ]
            if valid_candidates:
                return self.generate_and_save_memory(valid_candidates[0])

        cached = self.load_cached_memory()
        if cached is not None:
            return cached

        raise FileNotFoundError(
            f"No candidate memory file found at '{self.memory_path}' and no resume file provided. "
            "Please provide a resume file in config/candidate.local.yaml or via --resume <path>."
        )
